from docx import Document
import re
from openpyxl import Workbook

req_id_pattern = r"SYSREQ-\d+"

reqs_doc = Document("requirements.docx")
arc_doc = Document("architecture.docx")

source_reqs = []

for paragraph in reqs_doc.paragraphs:
    match = re.search(req_id_pattern, paragraph.text)
    if match:
        source_reqs.append(match.group())

# print(source_reqs)
print(f"Number of source requirements: {len(source_reqs)}")


def get_heading_number(paragraph):
    heading_regex = r"^((\d+\.)+) [\w\(\)\:\- ]+"
    # heading_regex = r"^((\d+\.)+).*"
    match = re.search(heading_regex, paragraph.text)
    if match:
        return match.group(0).strip()
    return None


arc_reqs = dict()
current_heading = ""

for paragraph in arc_doc.paragraphs:
    heading_number = get_heading_number(paragraph)
    if heading_number:
        current_heading = heading_number
        # print(current_heading)

    matches = re.findall(req_id_pattern, paragraph.text)
    if matches:
        for match in matches:
            arc_reqs[match] = arc_reqs.setdefault(match, []) + [current_heading]

# print(arc_reqs)

trace_table = dict()
untraced = []
for source_req in source_reqs:
    if source_req not in arc_reqs:
        untraced.append(source_req)
        trace_table[source_req] = ["Requirement not covered"]
    else:
        trace_table[source_req] = arc_reqs[source_req]

trace_table = dict(sorted(trace_table.items()))

# print(trace_table)
print(f"Number of untraced requirements: {len(untraced)}")
print(f"Untraced: {', '.join(untraced)}")

wb = Workbook()
ws = wb.active

ws.append(["Requirement UID", "Section in Architecture"])
for key, value in trace_table.items():
    ws.append([key, ", ".join(value)])

wb.save("trace.xlsx")
